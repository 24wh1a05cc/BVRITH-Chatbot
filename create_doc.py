"""
Script to create the BVRITH College FAQ knowledge base document (.docx)
Based on BVRIT Hyderabad College of Engineering for Women data.
Comprehensive reference for the BVRIT FAQ Chatbot · Academic Year 2025–26
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
import os


def create_bvrit_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('BVRIT HYDERABAD COLLEGE OF ENGINEERING FOR WOMEN', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('College Information Document')
    run.bold = True
    run.font.size = Pt(16)
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        'Comprehensive reference for the BVRIT FAQ Chatbot · Academic Year 2025–26\n'
        'Version 1.0 · Last Updated: 1 July 2025 · Document Owner: Office of the Dean (Academics)'
    )
    run.font.size = Pt(10)
    
    doc.add_paragraph('')
    
    # NOTICE
    doc.add_heading('NOTICE', level=2)
    doc.add_paragraph(
        'This document is generated for educational purposes as part of the GenAI & Agentic AI Engineering '
        'programme. It contains representative data modelled on typical engineering college structures in '
        'Telangana. Some figures, dates, and names are illustrative. Students should treat this as their '
        'authoritative grounding document for all chatbot exercises.'
    )
    
    doc.add_paragraph('')
    
    # ============================================================
    # SECTION 1: ABOUT BVRIT
    # ============================================================
    doc.add_heading('1. ABOUT BVRIT', level=1)
    
    doc.add_paragraph(
        'B V Raju Institute of Technology (BVRIT) Hyderabad College of Engineering for Women is an autonomous '
        'institution affiliated to Jawaharlal Nehru Technological University, Hyderabad (JNTUH). The college is '
        'located in Narsapur, Medak District, Telangana, on a 125-acre campus approximately 50 km from Hyderabad '
        'city centre.'
    )
    doc.add_paragraph(
        'BVRIT was established in 2009 by the Vishnu Educational Society, inspired by the vision of Sri B V Raju. '
        'The college exclusively admits women students and has grown to over 3,500 students across seven B.Tech programmes.'
    )
    
    doc.add_heading('1.1 Vision', level=2)
    doc.add_paragraph(
        'To be a premier institution empowering women through quality technical education, research, and innovation.'
    )
    
    doc.add_heading('1.2 Mission', level=2)
    doc.add_paragraph(
        'To provide industry-relevant engineering education with a focus on holistic development, ethical leadership, '
        'and social responsibility.'
    )
    
    doc.add_heading('1.3 Accreditations & Rankings', level=2)
    doc.add_paragraph(
        '• NAAC Accredited with \'A\' Grade\n'
        '• NBA Accredited: CSE, ECE, EEE, Mechanical, and IT programmes\n'
        '• Approved by AICTE, New Delhi\n'
        '• Autonomous status granted by UGC\n'
        '• NIRF 2024 Rank Band: 201–300 (Engineering category)'
    )
    
    # ============================================================
    # SECTION 2: DEPARTMENTS & PROGRAMMES
    # ============================================================
    doc.add_heading('2. DEPARTMENTS & PROGRAMMES', level=1)
    
    doc.add_paragraph(
        'BVRIT offers seven B.Tech programmes across seven departments. Each department maintains dedicated '
        'laboratories, a department library, and active industry partnerships.'
    )
    
    doc.add_heading('2.1 Computer Science & Engineering (CSE)', level=2)
    doc.add_paragraph(
        'The CSE department is the largest at BVRIT with 45 faculty members, including 8 with PhD qualifications '
        'and 12 pursuing doctoral research. The department operates 6 computing laboratories with 360 workstations. '
        'Key research areas include artificial intelligence, cybersecurity, cloud computing, and software engineering.'
    )
    doc.add_paragraph(
        'The department offers two specialisation tracks from the third year: Data Science and Cyber Security. '
        'Students choose one specialisation based on their interests and career goals. The CSE department has a '
        'dedicated placement cell that works with companies like TCS, Infosys, Wipro, Cognizant, Amazon, and Microsoft.'
    )
    doc.add_paragraph(
        'Notable achievements: 15 student research papers published in IEEE conferences in 2024–25. Three teams '
        'qualified for Smart India Hackathon 2025 national finals. The BVRIT ACM Student Chapter has 280 active members.'
    )
    
    doc.add_heading('2.2 Electronics & Communication Engineering (ECE)', level=2)
    doc.add_paragraph(
        'The ECE department has 38 faculty members, including 6 with PhD qualifications. The department maintains '
        '5 laboratories covering VLSI design, embedded systems, communication systems, signal processing, and IoT. '
        'Key research areas include antenna design, wireless communications, and MEMS technology.'
    )
    doc.add_paragraph(
        'The ECE department partners with Texas Instruments for an embedded systems lab and with National Instruments '
        'for a LabVIEW certification programme. Students have access to industry-standard EDA tools including Cadence '
        'Virtuoso and Synopsys Design Compiler.'
    )
    doc.add_paragraph(
        'Notable achievements: 8 student patents filed in 2024–25. The department\'s robotics team won second place '
        'at Robocon 2025 regional qualifiers.'
    )
    
    doc.add_heading('2.3 Electrical & Electronics Engineering (EEE)', level=2)
    doc.add_paragraph(
        'The EEE department has 28 faculty members, including 5 with PhD qualifications. The department operates '
        '4 laboratories: power electronics, electrical machines, control systems, and renewable energy. The department '
        'has a 10 kW rooftop solar installation used for teaching and research.'
    )
    doc.add_paragraph(
        'Key focus areas include smart grid technology, electric vehicle systems, and power system protection. '
        'The department collaborates with BHEL and NTPC for industrial training programmes.'
    )
    
    doc.add_heading('2.4 Mechanical Engineering', level=2)
    doc.add_paragraph(
        'The Mechanical Engineering department has 25 faculty members, including 4 with PhD qualifications. '
        'Facilities include workshops for manufacturing processes, a CAD/CAM laboratory, thermal engineering lab, '
        'and a material testing laboratory. The department has an MoU with Siemens for NX software training.'
    )
    doc.add_paragraph(
        'The department operates a student-run Formula SAE team that designs, builds, and races an open-wheel car '
        'annually. The team placed 8th nationally in Formula Bharat 2025.'
    )
    
    doc.add_heading('2.5 Information Technology (IT)', level=2)
    doc.add_paragraph(
        'The IT department has 32 faculty members, including 5 with PhD qualifications. The department focuses on '
        'web technologies, database systems, software testing, and network security. It operates 4 computing laboratories '
        'with 240 workstations.'
    )
    doc.add_paragraph(
        'The IT department runs an annual 24-hour hackathon (BVRIT HackIT) that attracts participants from 30+ colleges '
        'across Telangana. The department has active partnerships with Oracle Academy, Red Hat, and AWS Educate.'
    )
    
    doc.add_heading('2.6 CSE (Artificial Intelligence & Machine Learning)', level=2)
    doc.add_paragraph(
        'Established in 2021, the AI&ML programme has 18 faculty members, including 3 with PhD qualifications and 6 '
        'with industry experience in AI/ML. The department has a dedicated GPU computing lab with 8 NVIDIA A100 '
        'workstations for deep learning research.'
    )
    doc.add_paragraph(
        'Core specialisation areas: computer vision, natural language processing, reinforcement learning, and MLOps. '
        'Students complete a mandatory 6-month industry internship in the final year. Partner companies include Google, '
        'NVIDIA, and Qualcomm.'
    )
    
    doc.add_heading('2.7 CSE (Data Science)', level=2)
    doc.add_paragraph(
        'Established in 2021, the Data Science programme has 16 faculty members, including 2 with PhD qualifications '
        'and 5 with industry data science experience. The department emphasises statistical modelling, big data systems, '
        'and data engineering.'
    )
    doc.add_paragraph(
        'Students gain hands-on experience with tools including Apache Spark, Hadoop, Tableau, Power BI, and cloud '
        'platforms (AWS, Azure). The department partners with Deloitte and EY for case study projects.'
    )
    
    # ============================================================
    # SECTION 3: FEE STRUCTURE
    # ============================================================
    doc.add_heading('3. FEE STRUCTURE', level=1)
    
    doc.add_paragraph(
        'All fees are quoted in Indian Rupees (₹) for the academic year 2025–26. Fees are payable annually at the '
        'beginning of each academic year. The fee structure is approved by the Telangana State Council for Higher '
        'Education (TSCHE).'
    )
    
    doc.add_heading('3.1 Tuition Fees by Programme', level=2)
    doc.add_paragraph(
        'B.Tech CSE: ₹1,20,000 per year (₹4,80,000 total for 4 years)\n'
        'B.Tech CSE (AI&ML): ₹1,35,000 per year (₹5,40,000 total)\n'
        'B.Tech CSE (Data Science): ₹1,35,000 per year (₹5,40,000 total)\n'
        'B.Tech ECE: ₹1,10,000 per year (₹4,40,000 total)\n'
        'B.Tech EEE: ₹1,00,000 per year (₹4,00,000 total)\n'
        'B.Tech Mechanical: ₹1,00,000 per year (₹4,00,000 total)\n'
        'B.Tech IT: ₹1,10,000 per year (₹4,40,000 total)'
    )
    
    doc.add_heading('3.2 Hostel Fees', level=2)
    doc.add_paragraph(
        'Hostel accommodation is available for all students. The campus has 5 hostel blocks with a total capacity '
        'of 3,200 students.'
    )
    doc.add_paragraph(
        'Room Rent (Shared, 3-seater): ₹60,000 per year (₹2,40,000 total)\n'
        'Room Rent (Shared, 2-seater): ₹75,000 per year (₹3,00,000 total)\n'
        'Mess Charges (Vegetarian): ₹48,000 per year (₹1,92,000 total)\n'
        'Mess Charges (Non-Vegetarian): ₹54,000 per year (₹2,16,000 total)'
    )
    
    doc.add_heading('3.3 Other Fees', level=2)
    doc.add_paragraph(
        'Transport (College Bus): ₹45,000 per year (optional)\n'
        'Laboratory Fee: ₹15,000 per year\n'
        'Library & Digital Resources: ₹8,000 per year\n'
        'Examination Fee: ₹5,000 per semester (₹10,000/year)\n'
        'Student Activity Fee: ₹3,000 per year\n'
        'Caution Deposit (Refundable): ₹10,000 (one-time at admission)\n'
        'Admission Processing Fee: ₹5,000 (one-time at admission)'
    )
    
    doc.add_heading('3.4 Total Cost of Attendance (Estimated)', level=2)
    doc.add_paragraph(
        'For hostel residence with a 3-seater room and vegetarian mess:\n\n'
        'CSE: ₹10,71,000 (4 years)\n'
        'ECE: ₹10,31,000 (4 years)\n'
        'EEE: ₹9,91,000 (4 years)\n'
        'Mechanical: ₹9,91,000 (4 years)\n\n'
        'Note: Transport fees (₹45,000/year) are optional and not included in the totals above.'
    )
    
    # ============================================================
    # SECTION 4: SCHOLARSHIPS & FEE CONCESSIONS
    # ============================================================
    doc.add_heading('4. SCHOLARSHIPS & FEE CONCESSIONS', level=1)
    
    doc.add_paragraph(
        'BVRIT offers several scholarship schemes to support meritorious and economically disadvantaged students. '
        'Scholarships are applied as a percentage discount on the annual tuition fee only (not on hostel, mess, '
        'or other fees).'
    )
    
    doc.add_heading('4.1 Merit Scholarships', level=2)
    doc.add_paragraph(
        'Founder\'s Scholarship: EAMCET rank within top 1,000 — 100% tuition waiver (renewable if CGPA ≥ 8.5)\n'
        'Academic Excellence: EAMCET rank 1,001–5,000 — 50% tuition discount (renewable if CGPA ≥ 8.0)\n'
        'Merit Reward: EAMCET rank 5,001–15,000 — 25% tuition discount (renewable if CGPA ≥ 7.5)\n'
        'Sports Scholarship: National/state level sports achievement — 25% tuition discount (renewable)'
    )
    
    doc.add_heading('4.2 Need-Based Fee Concessions', level=2)
    doc.add_paragraph(
        'Economically Weaker Section (EWS): Family income below ₹8 lakh/year — 50% discount\n'
        'SC/ST Scholarship: Telangana state SC/ST students — Full tuition reimbursement via state govt\n'
        'BC Fee Reimbursement: Telangana BC students, income < ₹2 lakh — Full tuition reimbursement via state govt\n'
        'Sibling Discount: Second sibling currently enrolled at BVRIT — 10% discount on tuition\n\n'
        'Note: Government scholarships (SC/ST, BC) are subject to state government disbursement timelines. '
        'Students must pay the full fee at admission and receive reimbursement after government processing, '
        'which typically takes 3–6 months.'
    )
    
    # ============================================================
    # SECTION 5: ADMISSIONS
    # ============================================================
    doc.add_heading('5. ADMISSIONS', level=1)
    
    doc.add_heading('5.1 Eligibility', level=2)
    doc.add_paragraph(
        'Candidates must have passed the Intermediate examination (or equivalent) with Mathematics, Physics, and '
        'Chemistry (MPC) with a minimum of 45% aggregate marks (40% for reserved categories). A valid TS EAMCET '
        'or AP EAMCET rank is required for admission through counselling.'
    )
    doc.add_paragraph(
        'Direct admission under the management quota requires a minimum of 50% in MPC and a valid EAMCET rank. '
        'JEE Main qualified candidates are also eligible.'
    )
    
    doc.add_heading('5.2 Admission Process', level=2)
    doc.add_paragraph(
        '1. Obtain TS EAMCET or AP EAMCET rank.\n'
        '2. Participate in the TSCHE web counselling process and select BVRIT as a preference.\n'
        '3. Report to the college with original documents within 5 days of allotment.\n'
        '4. Complete document verification, fee payment, and hostel allocation.\n'
        '5. Attend the orientation programme on the first day of classes.'
    )
    
    doc.add_heading('5.3 Important Dates — Academic Year 2025–26', level=2)
    doc.add_paragraph(
        'TS EAMCET 2025 Exam: 15 May 2025 (Completed)\n'
        'TS EAMCET Results Announced: 5 June 2025 (Completed)\n'
        'TSCHE Web Counselling Round 1: 1 July – 15 July 2025 (Completed)\n'
        'TSCHE Web Counselling Round 2: 20 July – 31 July 2025 (Completed)\n'
        'Last Date for Round 2 Reporting: 5 August 2025 (Upcoming)\n'
        'Management Quota Admission Deadline: 15 August 2025 (Upcoming)\n'
        'Orientation Programme: 18 August 2025 (Upcoming)\n'
        'Classes Commence: 19 August 2025 (Upcoming)\n'
        'Last Date for Late Admission (with penalty): 31 August 2025 (Upcoming)\n'
        'Mid-Semester Exam 1: 15 October – 20 October 2025 (Upcoming)\n'
        'End-Semester Exam (Odd Semester): 1 December – 15 December 2025 (Upcoming)\n'
        'Even Semester Classes Begin: 5 January 2026 (Upcoming)\n'
        'Mid-Semester Exam 2: 15 March – 20 March 2026 (Upcoming)\n'
        'End-Semester Exam (Even Semester): 1 May – 15 May 2026 (Upcoming)'
    )
    doc.add_paragraph(
        'Note: All dates are subject to change based on TSCHE and JNTUH academic calendar notifications. '
        'Students should verify dates on the official BVRIT website (www.bvrithyderabad.edu.in) before making '
        'any decisions.'
    )
    
    # ============================================================
    # SECTION 6: PLACEMENTS
    # ============================================================
    doc.add_heading('6. PLACEMENTS', level=1)
    
    doc.add_paragraph(
        'The Training & Placement Cell coordinates campus recruitment activities, industry internships, and career '
        'development programmes. The placement season runs from August to April each year.'
    )
    
    doc.add_heading('6.1 Placement Statistics (2024–25)', level=2)
    doc.add_paragraph(
        'CSE: 205/228 placed (89.9%), Highest ₹24.0 LPA, Average ₹6.2 LPA\n'
        'ECE: 132/165 placed (80.0%), Highest ₹18.5 LPA, Average ₹5.1 LPA\n'
        'EEE: 78/108 placed (72.2%), Highest ₹12.0 LPA, Average ₹4.3 LPA\n'
        'Mechanical: 35/52 placed (67.3%), Highest ₹10.5 LPA, Average ₹4.0 LPA\n'
        'IT: 147/168 placed (87.5%), Highest ₹21.0 LPA, Average ₹5.8 LPA\n'
        'AI&ML: 89/105 placed (84.8%), Highest ₹22.0 LPA, Average ₹6.5 LPA\n'
        'Data Science: 85/102 placed (83.3%), Highest ₹20.0 LPA, Average ₹6.0 LPA'
    )
    
    doc.add_heading('6.2 Top Recruiters (2024–25)', level=2)
    doc.add_paragraph(
        'IT Services: TCS, Infosys, Wipro, Cognizant, HCL Technologies, Tech Mahindra, Capgemini\n'
        'Product Companies: Amazon, Microsoft, Google (internships), Salesforce, ServiceNow\n'
        'Core Companies: BHEL, L&T, Tata Motors, Schneider Electric, Siemens\n'
        'Startups & Mid-size: Zoho, Freshworks, Razorpay, PhonePe, Swiggy (tech roles)\n'
        'Consulting: Deloitte, EY, KPMG (technology consulting roles)'
    )
    
    doc.add_heading('6.3 Internship Programme', level=2)
    doc.add_paragraph(
        'All students complete a mandatory internship: 6 weeks after Year 3 (summer) and 6 months in Year 4 '
        '(for select programmes). Stipends range from ₹10,000 to ₹60,000 per month depending on the company and '
        'role. In 2024–25, 78% of students who completed internships at product companies received pre-placement offers.'
    )
    
    doc.add_paragraph(
        '⚠️ Disclaimer: Placement statistics are historical data from the 2024–25 academic year and do not '
        'guarantee future outcomes. Individual placement depends on academic performance, skills, interview '
        'performance, and market conditions. BVRIT does not guarantee placement to any student.'
    )
    
    # ============================================================
    # SECTION 7: CAMPUS & FACILITIES
    # ============================================================
    doc.add_heading('7. CAMPUS & FACILITIES', level=1)
    
    doc.add_heading('7.1 Academic Facilities', level=2)
    doc.add_paragraph(
        'Central Library: 50,000+ volumes, 2,500 e-journals, 24/7 digital access, 200-seat reading hall\n'
        'Computing Labs: 1,200+ workstations across all departments, 100 Mbps dedicated internet\n'
        'Smart Classrooms: all classrooms equipped with interactive displays and lecture recording\n'
        'Research Centre: shared facility with high-performance computing, GPU cluster, and 3D printing\n'
        'Language Lab: 60-seat lab for English communication skills and soft skills training'
    )
    
    doc.add_heading('7.2 Hostel Facilities', level=2)
    doc.add_paragraph(
        'BVRIT provides on-campus hostel accommodation for all students. The campus has 5 hostel blocks (Jasmine, '
        'Lotus, Orchid, Rose, and Tulip) with a total capacity of 3,200 students. Each room is furnished with beds, '
        'study tables, chairs, and wardrobes.'
    )
    doc.add_paragraph(
        'Hostel facilities include: 24/7 Wi-Fi (50 Mbps per hostel block), RO purified drinking water, hot water '
        'supply (6–8 AM, 6–8 PM), laundry service (outsourced, ₹300/month), common rooms with TV and indoor games, '
        'and 24-hour security with CCTV surveillance.'
    )
    doc.add_paragraph(
        'The mess serves three meals daily (breakfast, lunch, dinner) plus evening snacks. The menu rotates weekly '
        'and accommodates vegetarian and non-vegetarian preferences. A student mess committee reviews food quality monthly.'
    )
    
    doc.add_heading('7.3 Sports & Recreation', level=2)
    doc.add_paragraph(
        'Indoor: badminton courts (4), table tennis (6 tables), chess room, yoga hall, gymnasium\n'
        'Outdoor: cricket ground, football field, basketball courts (2), volleyball courts (2), tennis court, 400m athletic track\n'
        'Annual sports meet: BVRIT Olympia (held in February)\n'
        'Professional coaching available for: badminton, basketball, and athletics'
    )
    
    doc.add_heading('7.4 Transport', level=2)
    doc.add_paragraph(
        'BVRIT operates 35 college buses covering routes across Hyderabad, Secunderabad, Kukatpally, Miyapur, '
        'Ameerpet, Dilsukhnagar, and LB Nagar. Buses depart from designated pickup points at 7:00 AM and return '
        'by 5:30 PM. The annual transport fee is ₹45,000. Students using their own transport (two-wheelers are '
        'permitted with valid licence and helmet) do not pay this fee. Parking is available on campus at no additional cost.'
    )
    
    doc.add_heading('7.5 Medical & Wellness', level=2)
    doc.add_paragraph(
        'A full-time medical officer is available on campus (9 AM – 5 PM, Monday–Saturday). Emergency medical support '
        'is available 24/7. The nearest hospital is Apollo Hospitals, Jubilee Hills (45 km). An ambulance is stationed '
        'on campus. All students are covered under a group medical insurance policy (₹1 lakh coverage) included in the '
        'admission fee.'
    )
    
    # ============================================================
    # SECTION 8: KEY FACULTY
    # ============================================================
    doc.add_heading('8. KEY FACULTY', level=1)
    
    doc.add_heading('8.1 College Leadership', level=2)
    doc.add_paragraph(
        'Principal: Dr. K. Lakshmi Prasad, PhD (IIT Madras), 28 years of academic experience\n'
        'Vice-Principal: Dr. M. Sravanthi, PhD (NIT Warangal), 22 years in academia\n'
        'Dean (Academics): Dr. P. Raghavendra, PhD (JNTUH), responsible for curriculum and quality\n'
        'Dean (Student Affairs): Dr. S. Kavitha, PhD (Osmania University), responsible for student welfare\n'
        'Training & Placement Officer: Prof. N. Venkat Reddy, 18 years of industry + academic experience'
    )
    
    doc.add_heading('8.2 Department Heads', level=2)
    doc.add_paragraph(
        'CSE: Dr. A. Padmavathi, PhD (IIT Hyderabad) — Machine Learning, NLP\n'
        'CSE (AI&ML): Dr. R. Swathi, PhD (IIIT Hyderabad) — Deep Learning, Computer Vision\n'
        'CSE (DS): Dr. B. Lakshmi, PhD (IISc Bangalore) — Big Data Analytics, Statistical Learning\n'
        'ECE: Dr. K. Anitha, PhD (NIT Warangal) — VLSI Design, Embedded Systems\n'
        'EEE: Dr. T. Manohar, PhD (IIT Bombay) — Power Electronics, Smart Grids\n'
        'Mechanical: Dr. V. Anuradha, PhD (NIT Trichy) — CAD/CAM, Additive Manufacturing\n'
        'IT: Dr. S. Priya, PhD (JNTUH) — Cloud Computing, Network Security'
    )
    
    # ============================================================
    # SECTION 9: STUDENT SUPPORT SERVICES
    # ============================================================
    doc.add_heading('9. STUDENT SUPPORT SERVICES', level=1)
    
    doc.add_heading('9.1 Student Counselling Centre', level=2)
    doc.add_paragraph(
        'The Student Counselling Centre provides free, confidential support for academic stress, personal issues, '
        'and mental health concerns. Two full-time counsellors are available Monday–Saturday, 9 AM – 5 PM. '
        'Appointments can be booked in person at the Admin Block, Room 105, or via email at counselling@bvrit.ac.in. '
        'Walk-ins are welcome.'
    )
    doc.add_paragraph(
        'In case of crisis: Students in distress can contact the counselling centre directly or call the 24-hour '
        'student helpline at 08455-221144. For after-hours emergencies, contact hostel wardens or campus security '
        '(available 24/7).'
    )
    doc.add_paragraph(
        'External crisis resources: iCall — 9152987821 (Mon–Sat, 8 AM – 10 PM). Vandrevala Foundation Helpline — '
        '1860-2662-345 (24/7). NIMHANS Helpline — 080-46110007.'
    )
    
    doc.add_heading('9.2 Anti-Ragging Committee', level=2)
    doc.add_paragraph(
        'BVRIT has a zero-tolerance policy on ragging. The Anti-Ragging Committee can be contacted at '
        'antiragging@bvrit.ac.in or the UGC helpline 1800-180-5522. All complaints are investigated within 48 hours.'
    )
    
    doc.add_heading('9.3 Grievance Redressal', level=2)
    doc.add_paragraph(
        'Students can submit grievances through the online portal (portal.bvrit.ac.in/grievance) or in writing to '
        'the Dean (Student Affairs). All grievances receive a response within 7 working days. The Grievance Redressal '
        'Committee meets monthly to review pending cases.'
    )
    
    # ============================================================
    # SECTION 10: CONTACT INFORMATION
    # ============================================================
    doc.add_heading('10. CONTACT INFORMATION', level=1)
    
    doc.add_paragraph(
        'General Enquiries: info@bvrit.ac.in — 08455-221100\n'
        'Admissions Office: admissions@bvrit.ac.in — 08455-221111\n'
        'Fee Payment Queries: accounts@bvrit.ac.in — 08455-221122\n'
        'Training & Placement: placements@bvrit.ac.in — 08455-221133\n'
        'Hostel & Accommodation: hostel@bvrit.ac.in — 08455-221144\n'
        'Student Counselling: counselling@bvrit.ac.in — 08455-221155\n'
        'Transport: transport@bvrit.ac.in — 08455-221166\n'
        'Examination Cell: exams@bvrit.ac.in — 08455-221177\n'
        'Principal\'s Office: principal@bvrit.ac.in — 08455-221100 (ext. 101)\n'
        'Dean (Student Affairs): dean.sa@bvrit.ac.in — 08455-221100 (ext. 105)'
    )
    
    doc.add_heading('10.1 Campus Address', level=2)
    doc.add_paragraph(
        'B V Raju Institute of Technology (BVRIT)\n'
        'Hyderabad College of Engineering for Women\n'
        'Narsapur, Medak District\n'
        'Telangana — 502313\n'
        'India'
    )
    doc.add_paragraph('Website: www.bvrithyderabad.edu.in')
    
    # Save the document
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'bvrit_college_info.docx')
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    print(f"Document created successfully with 10 main sections.")


if __name__ == '__main__':
    create_bvrit_document()